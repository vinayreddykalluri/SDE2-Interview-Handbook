#!/usr/bin/env python3
"""Assemble Markdown and build the print PDF and DOCX editions."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import math
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    CondPageBreak,
    Flowable,
    Frame,
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    LongTable,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    XPreformatted,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.platypus.doctemplate import ActionFlowable


ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT / "content" / "master"
APPENDICES = BOOK / "appendices"
BUILD = ROOT / "build"
TMP = ROOT / "tmp" / "pdfs"
DIST = ROOT / "dist"

MASTER_MD = DIST / "00-start-here" / "java-sde2-interview-book.md"
FINAL_DOCX = DIST / "00-start-here" / "java-sde2-interview-book.docx"
FINAL_PDF = DIST / "00-start-here" / "java-sde2-interview-book.pdf"
REFERENCE_DOCX = BUILD / "reference.docx"

TITLE = "Java Foundations to Advanced Engineering"
SUBTITLE = "A Complete SDE-2 Interview Preparation Guide"
EDITION = "Java 21 baseline - Java 17 and Java 21 features - 2026 edition"
AUTHOR = "Vinay Reddy Kalluri"

NAVY = HexColor("#102A43")
BLUE = HexColor("#176B87")
DARK_BLUE = HexColor("#164E63")
CYAN = HexColor("#237A78")
GOLD = HexColor("#C57A00")
GOLD_TEXT = HexColor("#8A5A00")
GREEN = HexColor("#247A65")
RED = HexColor("#A33A3A")
INK = HexColor("#1F2933")
MUTED = HexColor("#52616F")
PALE = HexColor("#E8F1F5")
LIGHT = HexColor("#F5F7FA")
LINE = HexColor("#B9C7D3")
PATTERN_BG = HexColor("#E8F5F3")
GUIDE_BG = HexColor("#FFF7E6")
TABLE_ALT = HexColor("#F4F8FA")
PRACTICE_BG = HexColor("#EAF6EF")
WARNING_BG = HexColor("#FFF1E8")
SUMMARY_BG = HexColor("#EEF2F6")

PAGE_W, PAGE_H = letter
MARGIN = 58
CONTENT_W = PAGE_W - 2 * MARGIN
COVER_W = 6.5 * inch
FRAME_H = PAGE_H - 62 - 56


CHAPTERS: list[tuple[int, str]] = [
    (1, "Why Java Exists"),
    (2, "JDK, JRE, JVM, Editions, and Distributions"),
    (3, "Compilation, Bytecode, and the Execution Pipeline"),
    (4, "JVM Architecture"),
    (5, "Class Loading, Linking, and Initialization"),
    (6, "Runtime Data Areas"),
    (7, "Object Creation and Memory Layout"),
    (8, "Java Stacks, Method Calls, and Recursion"),
    (9, "Garbage Collection"),
    (10, "Execution Engine, JIT Compilation, and Safepoints"),
    (11, "The Java Memory Model from First Principles"),
    (12, "Variables, Primitive Types, Literals, and Numeric Semantics"),
    (13, "Operators, Expressions, and Control Flow"),
    (14, "Methods, Overloading, Varargs, and Pass-by-Value"),
    (15, "Arrays, Strings, Text Blocks, and Unicode"),
    (16, "Classes, Objects, Access Control, and Packages"),
    (17, "Inheritance, Polymorphism, and Composition"),
    (18, "Interfaces, Abstract Classes, Sealed Types, and Pattern Matching"),
    (19, "Equality, Hashing, Immutability, and Records"),
    (20, "Exceptions and Resource Management"),
    (21, "Nested Types, Enums, Annotations, and Reflection"),
    (22, "Generics, Variance, Type Erasure, and Heap Pollution"),
    (23, "Lambdas, Method References, and Functional Interfaces"),
    (24, "Java 17 and Java 21 Language and Platform Features"),
    (25, "Collections Framework Architecture"),
    (26, "ArrayList, LinkedList, and List Trade-offs"),
    (27, "HashMap, HashSet, and Hashing Internals"),
    (28, "TreeMap, TreeSet, Ordering, and Navigable Collections"),
    (29, "Queues, Deques, PriorityQueue, and Heaps"),
    (30, "Comparable, Comparator, Sorting, and Selection"),
    (31, "Streams, Collectors, Optional, and Spliterators"),
    (32, "Java I/O, NIO.2, Files, Buffers, and Serialization Boundaries"),
    (33, "Threads, Lifecycle, Interruption, and Cancellation"),
    (34, "Synchronization, Intrinsic Locks, and Explicit Locks"),
    (35, "Volatile, Atomics, CAS, and Happens-Before in Practice"),
    (36, "Executors, Futures, CompletableFuture, and Work Scheduling"),
    (37, "Concurrent Collections and Virtual Threads"),
    (38, "Concurrency Failure Modes, Testing, and Design Patterns"),
    (39, "Performance Methodology and JMH Benchmarking"),
    (40, "JVM Diagnostics with jcmd, jstack, jmap, JFR, and Mission Control"),
    (41, "Memory Leaks, GC Incidents, and Tuning Playbooks"),
    (42, "Complexity and the SDE-2 Problem-Solving Method"),
    (43, "Arrays, Strings, Hashing, Two Pointers, Sliding Windows, and Prefix Sums"),
    (44, "Linked Lists, Stacks, Queues, and Monotonic Structures"),
    (45, "Trees, BSTs, Heaps, and Tries"),
    (46, "Graphs, Topological Sort, Shortest Paths, and Union-Find"),
    (47, "Recursion, Backtracking, Greedy Reasoning, and Dynamic Programming"),
    (48, "The Java Coding Interview Playbook"),
    (49, "Clean Java APIs, SOLID Design, and Low-Level Design Patterns"),
    (50, "Backend Java Boundaries: JDBC, Transactions, Serialization, and Services"),
    (51, "Testing, Build Tools, Static Analysis, and Dependency Management"),
    (52, "Secure and Reliable Java"),
    (53, "SDE-2 Java Interview Question Bank"),
    (54, "Eight-Week Study Plan and Mock Interview Loops"),
]

PARTS: list[tuple[str, str, int, int]] = [
    ("Part I", "Java and the Computing Model", 1, 3),
    ("Part II", "JVM Architecture and Memory", 4, 11),
    ("Part III", "Java Language Engineering", 12, 24),
    ("Part IV", "Collections, Streams, and I/O", 25, 32),
    ("Part V", "Concurrency and Multithreading", 33, 38),
    ("Part VI", "Performance, Diagnostics, and Reliability", 39, 41),
    ("Part VII", "Data Structures and Algorithms in Java", 42, 48),
    ("Part VIII", "Engineering Practice and Interview Readiness", 49, 54),
]

APPENDIX_MAP: list[tuple[str, str, str]] = [
    ("A", "Java Syntax and Language Quick Reference", "a-java-quick-reference.md"),
    ("B", "Collection Complexity and Selection Matrix", "b-collection-complexity.md"),
    ("C", "JVM Tools, Flags, and Incident Commands", "c-jvm-tools-flags.md"),
    ("D", "Java 17 and Java 21 Feature Matrix", "d-java-feature-matrix.md"),
    ("E", "Exercise Hints and Selected Solutions", "e-exercise-solutions.md"),
    ("F", "Glossary", "f-glossary.md"),
    ("G", "Primary References and Further Reading", "g-references.md"),
]

FIGURES: dict[int, tuple[str, str]] = {
    3: ("01-compilation-pipeline.png", "Compilation and execution from source to native instructions"),
    4: ("02-jvm-architecture.png", "JVM architecture and representative HotSpot runtime components"),
    5: ("03-class-lifecycle.png", "Loading, linking, and initialization lifecycle"),
    6: ("04-runtime-data-areas.png", "Shared and per-thread runtime data areas"),
    7: ("05-object-layout.png", "Object creation semantics and representative HotSpot layout"),
    9: ("06-gc-reachability.png", "Reachability and the generational collection heuristic"),
    11: ("07-happens-before.png", "A happens-before chain through volatile synchronization"),
    25: ("08-collections-map.png", "Collections framework semantic map"),
    27: ("09-hashmap-put.png", "Representative HashMap insertion path"),
    36: ("10-executor-model.png", "Executor submission, queueing, work, and completion"),
    37: ("11-virtual-threads.png", "Virtual threads mounted on carrier threads"),
    48: ("12-interview-loop.png", "Repeatable coding interview control loop"),
}

ASCII_REPLACEMENTS = str.maketrans(
    {
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2212": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",
    }
)


def ascii_safe(text: str) -> str:
    return text.translate(ASCII_REPLACEMENTS)


def slugify(text: str) -> str:
    value = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return value or "section"


def chapter_file(number: int, allow_incomplete: bool) -> Path | None:
    matches = sorted(BOOK.glob(f"{number:02d}-*.md"))
    if len(matches) == 1:
        return matches[0]
    if len(matches) > 1:
        raise RuntimeError(f"Multiple source files found for Chapter {number}: {matches}")
    if allow_incomplete:
        return None
    raise RuntimeError(f"Missing source file for Chapter {number} under {BOOK}")


def strip_first_h1(text: str) -> str:
    lines = text.splitlines()
    for i, line in enumerate(lines):
        if re.match(r"^#\s+", line):
            del lines[i]
            if i < len(lines) and not lines[i].strip():
                del lines[i]
            break
    return "\n".join(lines).strip()


def number_sections(text: str, prefix: str) -> str:
    counter = 0
    output: list[str] = []
    for line in text.splitlines():
        match = re.match(r"^(##)\s+(.*)$", line)
        if match:
            counter += 1
            title = re.sub(r"^(?:\d+(?:\.\d+)*|[A-Z]\.\d+)\s*[.:-]?\s*", "", match.group(2)).strip()
            output.append(f"## {prefix}.{counter} {title}")
        else:
            output.append(line)
    return "\n".join(output)


def front_text() -> str:
    files = [
        BOOK / "00-start-here.md",
        BOOK / "00-preface.md",
        BOOK / "00-about-the-author.md",
        BOOK / "00-how-to-use-this-book.md",
        BOOK / "00-study-roadmap.md",
    ]
    missing = [p for p in files if not p.exists()]
    if missing:
        raise RuntimeError(f"Missing front matter: {missing}")
    return "\n\n".join(ascii_safe(p.read_text(encoding="utf-8").strip()) for p in files)


def copyright_page() -> str:
    return ascii_safe(
        """# Copyright and Disclaimer

Copyright 2026 Vinay Reddy Kalluri and credited contributors. Open educational edition.

Book prose, exercises, diagrams, and published editions are licensed under Creative Commons Attribution 4.0 International (CC BY 4.0). Build scripts and source code are licensed under MIT. Individual credit is recorded in AUTHORS.md and Git history.

Repository: https://github.com/vinayreddykalluri/SDE2-Interview-Handbook. Attribution does not imply endorsement. Java and related marks are owned by their respective holders. Company names are used only to describe common interview markets.

The Java platform, JVM implementations, tools, flags, support policies, and licensing terms evolve. Examples target Java 21 unless a section states otherwise. Verify release-specific and vendor-specific behavior against the primary sources in Appendix G before using it in production.

The examples are provided without warranty. Review correctness, security, performance, operational, and legal requirements before adapting any example to a production system.
"""
    )


def detailed_contents() -> str:
    lines = ["# Contents", "", "## Front Matter", "", "- Copyright and Disclaimer", "- Start Here - Choose Your Route", "- Preface", "- About the Author", "- How to Use This Book", "- Study Roadmap", ""]
    title_by_number = dict(CHAPTERS)
    for part_label, part_title, start, end in PARTS:
        lines.extend([f"## {part_label} - {part_title}", ""])
        for number in range(start, end + 1):
            title = title_by_number[number]
            anchor = slugify(f"chapter-{number}-{title}")
            lines.append(f"- [Chapter {number} - {title}](#{anchor})")
        lines.append("")
    lines.extend(["## Appendices", ""])
    for letter, title, _ in APPENDIX_MAP:
        lines.append(f"- [Appendix {letter} - {title}](#{slugify(f'appendix-{letter}-{title}')})")
    return "\n".join(lines)


def assemble_sources(allow_incomplete: bool = False) -> tuple[str, str]:
    BUILD.mkdir(parents=True, exist_ok=True)
    title_by_number = dict(CHAPTERS)
    pieces: list[str] = []
    content_pieces: list[str] = []

    metadata = f"""---
title: "{TITLE}"
subtitle: "{SUBTITLE}"
author: "{AUTHOR}"
date: "{EDITION}"
lang: en-US
rights: "Copyright 2026 {AUTHOR} and credited contributors. CC BY 4.0 content; MIT code."
---
"""
    pieces.extend([metadata.strip(), copyright_page(), detailed_contents(), front_text()])
    content_pieces.append(front_text())

    for part_label, part_title, start, end in PARTS:
        part_heading = f"# {part_label} - {part_title}"
        pieces.append(part_heading)
        content_pieces.append(part_heading)
        for number in range(start, end + 1):
            path = chapter_file(number, allow_incomplete)
            if path is None:
                chapter_body = "## Learning objectives\n\nThis chapter is pending."
            else:
                chapter_body = strip_first_h1(ascii_safe(path.read_text(encoding="utf-8")))
            chapter_body = number_sections(chapter_body, str(number))
            heading = f"# Chapter {number} - {title_by_number[number]}"
            figure = ""
            if number in FIGURES:
                filename, caption = FIGURES[number]
                figure = f"\n\n![Figure {number}.1 - {caption}](assets/diagrams/{filename})\n"
            block = f"{heading}{figure}\n\n{chapter_body}".strip()
            pieces.append(block)
            content_pieces.append(block)

    appendix_part = "# Appendices"
    pieces.append(appendix_part)
    content_pieces.append(appendix_part)
    for letter, title, filename in APPENDIX_MAP:
        path = APPENDICES / filename
        if not path.exists():
            if allow_incomplete:
                body = "## Scope\n\nThis appendix is pending."
            else:
                raise RuntimeError(f"Missing appendix source: {path}")
        else:
            body = strip_first_h1(ascii_safe(path.read_text(encoding="utf-8")))
        body = number_sections(body, letter)
        block = f"# Appendix {letter} - {title}\n\n{body}"
        pieces.append(block)
        content_pieces.append(block)

    master = "\n\n<!-- PAGE_BREAK -->\n\n".join(pieces).strip() + "\n"
    content = "\n\n".join(content_pieces).strip() + "\n"
    MASTER_MD.write_text(master, encoding="utf-8")
    content_path = BUILD / "book-content.md"
    content_path.write_text(content, encoding="utf-8")
    docx_input = BUILD / "book-docx.md"
    docx_input.write_text(
        "\n\n".join(
            [metadata.strip(), copyright_page(), detailed_contents(), front_text(), *content_pieces[1:]]
        ).strip()
        + "\n",
        encoding="utf-8",
    )
    return str(content_path), str(docx_input)


def run(cmd: Sequence[str], cwd: Path = ROOT, capture: bool = False) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        list(map(str, cmd)),
        cwd=str(cwd),
        check=True,
        text=True,
        capture_output=capture,
    )


def ensure_reference_doc() -> None:
    script = ROOT / "scripts" / "create_reference_doc.py"
    run([sys.executable, script])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "B8C4D1")


def table_widths(table) -> list[int]:
    cols = len(table.columns)
    if cols == 0:
        return []
    weights: list[float] = []
    for col_idx in range(cols):
        longest = 1
        for row in table.rows:
            if col_idx < len(row.cells):
                text = " ".join(p.text for p in row.cells[col_idx].paragraphs)
                longest = max(longest, min(len(text), 90))
        weights.append(max(3.0, math.sqrt(longest)))
    total = sum(weights)
    widths = [max(720, round(9360 * weight / total)) for weight in weights]
    delta = 9360 - sum(widths)
    widths[-1] += delta
    if widths[-1] < 720:
        take = 720 - widths[-1]
        widths[-1] = 720
        widths[widths.index(max(widths[:-1]))] -= take
    return widths


def apply_table_geometry(table) -> None:
    widths = table_widths(table)
    if not widths:
        return
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_idx, cell in enumerate(row.cells):
            width = widths[min(col_idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
                    if row_idx == 0:
                        run.bold = True
    set_table_borders(table)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def remove_table_borders(table, bottom: bool = False) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    node = OxmlElement("w:bottom")
    node.set(qn("w:val"), "single" if bottom else "nil")
    if bottom:
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), "B8C4D1")
    borders.append(node)


def add_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        base = header.paragraphs[0]
        clear_paragraph(base)
        base.paragraph_format.space_after = Pt(0)
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.autofit = False
        widths = [4680, 4680]
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            node = OxmlElement("w:gridCol")
            node.set(qn("w:w"), str(width))
            grid.append(node)
        for idx, cell in enumerate(table.rows[0].cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            if idx == 0:
                run = p.add_run(TITLE)
            else:
                add_field(p, ' STYLEREF "Heading 1" ', "SDE-2 Interview Guide")
                run = p.runs[-1]
            run.font.name = "Calibri"
            run.font.size = Pt(8.5)
        remove_table_borders(table, bottom=True)

        footer = section.footer
        p = footer.paragraphs[0]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("Java SDE-2 Guide  |  ")
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        add_field(p, " PAGE ", "1")
        for item in p.runs:
            item.font.name = "Calibri"
            item.font.size = Pt(8.5)


def patch_docx_xml(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        files = {name: source.read(name) for name in source.namelist()}
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    numbering_name = "word/numbering.xml"
    if numbering_name in files:
        root = etree.fromstring(files[numbering_name])
        for lvl in root.xpath(".//w:lvl", namespaces=ns):
            ilvl = int(lvl.get(qn("w:ilvl"), "0"))
            left = 540 + ilvl * 360
            ppr = lvl.find(qn("w:pPr"))
            if ppr is None:
                ppr = etree.SubElement(lvl, qn("w:pPr"))
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = etree.SubElement(ppr, qn("w:ind"))
            ind.set(qn("w:left"), str(left))
            ind.set(qn("w:hanging"), "270")
            tabs = ppr.find(qn("w:tabs"))
            if tabs is None:
                tabs = etree.SubElement(ppr, qn("w:tabs"))
            for child in list(tabs):
                tabs.remove(child)
            tab = etree.SubElement(tabs, qn("w:tab"))
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(left))
            spacing = ppr.find(qn("w:spacing"))
            if spacing is None:
                spacing = etree.SubElement(ppr, qn("w:spacing"))
            spacing.set(qn("w:after"), "80")
            spacing.set(qn("w:line"), "300")
            spacing.set(qn("w:lineRule"), "auto")
        files[numbering_name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    settings_name = "word/settings.xml"
    if settings_name in files:
        root = etree.fromstring(files[settings_name])
        update = root.find(qn("w:updateFields"))
        if update is None:
            update = etree.SubElement(root, qn("w:updateFields"))
        update.set(qn("w:val"), "true")
        files[settings_name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")

    tmp_path = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name, data in files.items():
            target.writestr(name, data)
    tmp_path.replace(path)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_rule(paragraph, color: str, size: str = "8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    for edge in ("top", "bottom"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), color)
        borders.append(node)


def format_cover_paragraph(
    paragraph,
    *,
    size: float,
    color: str,
    bold: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.widow_control = True
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)
        rfonts = run._r.get_or_add_rPr().get_or_add_rFonts()
        for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            rfonts.attrib.pop(qn(f"w:{attribute}"), None)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attribute}"), "Arial")


def enhance_docx_cover(doc: Document) -> None:
    title = next((p for p in doc.paragraphs if p.style and p.style.name == "Title"), None)
    subtitle = next((p for p in doc.paragraphs if p.style and p.style.name == "Subtitle"), None)
    author = next((p for p in doc.paragraphs if p.style and p.style.name == "Author"), None)
    date = next((p for p in doc.paragraphs if p.style and p.style.name == "Date"), None)
    copyright_heading = next((p for p in doc.paragraphs if p.text.strip() == "Copyright and Disclaimer"), None)
    if any(item is None for item in (title, subtitle, author, date, copyright_heading)):
        raise RuntimeError("DOCX title-page metadata paragraphs were not generated as expected")

    banner = title.insert_paragraph_before("JAVA | JVM | CONCURRENCY | SYSTEM DESIGN")
    format_cover_paragraph(banner, size=9.5, color="FFFFFF", bold=True, after=28, line_spacing=1.35)
    set_paragraph_shading(banner, "0B2545")

    kicker = title.insert_paragraph_before("JAVA ENGINEERING INTERVIEW FIELD GUIDE")
    format_cover_paragraph(kicker, size=9.5, color="C58A22", bold=True, after=12)

    format_cover_paragraph(title, size=31, color="0B2545", bold=True, after=10, line_spacing=1.0)
    format_cover_paragraph(subtitle, size=15, color="1F4D78", after=11, line_spacing=1.1)

    author_label = author.insert_paragraph_before("WRITTEN BY")
    format_cover_paragraph(author_label, size=8.5, color="C58A22", bold=True, before=5, after=2)
    format_cover_paragraph(author, size=14, color="0B2545", bold=True, after=18)

    deck = doc.add_paragraph(
        "A first-principles guide to Java semantics, JVM internals, collections, concurrency, performance, algorithms, and production engineering."
    )
    format_cover_paragraph(deck, size=10.2, color="52606D", after=16, line_spacing=1.15)
    deck.paragraph_format.left_indent = Inches(0.35)
    deck.paragraph_format.right_indent = Inches(0.35)
    date._p.addprevious(deck._p)

    scope = doc.add_paragraph("LANGUAGE  |  JVM  |  COLLECTIONS  |  CONCURRENCY  |  DSA  |  BACKEND")
    format_cover_paragraph(scope, size=8.7, color="0B2545", bold=True, after=7, line_spacing=1.25)
    set_paragraph_shading(scope, "F2F4F7")
    set_paragraph_rule(scope, "C58A22", "6")
    date._p.addprevious(scope._p)

    stats = doc.add_paragraph("54 CHAPTERS  |  7 APPENDICES  |  JAVA 21 BASELINE  |  8-WEEK STUDY PLAN")
    format_cover_paragraph(stats, size=8.3, color="52606D", bold=True, before=9, after=7, line_spacing=1.15)
    date._p.addprevious(stats._p)

    format_cover_paragraph(date, size=8.8, color="52606D", before=18, after=2)
    date.paragraph_format.keep_with_next = True
    edition_note = doc.add_paragraph("Independent printable interview study edition")
    format_cover_paragraph(edition_note, size=8.5, color="52606D", after=0)
    edition_note.paragraph_format.keep_with_next = False
    copyright_heading._p.addprevious(edition_note._p)


def postprocess_docx(path: Path) -> None:
    doc = Document(path)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Part ") or text == "Appendices":
            try:
                paragraph.style = doc.styles["Book Part"]
            except KeyError:
                pass
        if paragraph.style and paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.widow_control = True

    # Pandoc's native Table and Compact styles are intentionally preserved.
    # Rewriting tblPr/tcPr children with python-docx can create a technically
    # out-of-order OOXML tree that Word tolerates but LibreOffice misrenders.
    enhance_docx_cover(doc)
    add_headers_and_footers(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = SUBTITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.comments = "Generated from the chapter Markdown sources."
    doc.save(path)


def build_docx(input_path: str) -> None:
    ensure_reference_doc()
    if shutil.which("pandoc") is None:
        raise RuntimeError("pandoc is required to build the DOCX")
    cmd = [
        "pandoc",
        input_path,
        "--from=gfm+yaml_metadata_block+pipe_tables+task_lists",
        "--to=docx",
        "--standalone",
        "--syntax-highlighting=tango",
        f"--resource-path={ROOT}",
        f"--reference-doc={REFERENCE_DOCX}",
        "-o",
        FINAL_DOCX,
    ]
    run(cmd)
    postprocess_docx(FINAL_DOCX)
    out = ROOT / "output" / "docx"
    out.mkdir(parents=True, exist_ok=True)
    shutil.copy2(FINAL_DOCX, out / FINAL_DOCX.name)


def register_fonts() -> dict[str, str]:
    # The release uses a book-oriented serif for sustained reading, a clean
    # sans serif for navigation/headings, and a high-legibility mono for code.
    # Every role has open-font fallbacks so rebuilding is not macOS-specific.
    candidates: dict[str, list[tuple[str, int | None]]] = {
        "BookBody": [
            ("/System/Library/Fonts/Supplemental/Charter.ttc", 0),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf", None),
        ],
        "BookBody-Bold": [
            ("/System/Library/Fonts/Supplemental/Charter.ttc", 3),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf", None),
        ],
        "BookBody-Italic": [
            ("/System/Library/Fonts/Supplemental/Charter.ttc", 1),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf", None),
        ],
        "BookSans": [
            ("/System/Library/Fonts/Avenir Next.ttc", 7),
            ("/System/Library/Fonts/Supplemental/Arial.ttf", None),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", None),
        ],
        "BookSans-Bold": [
            ("/System/Library/Fonts/Avenir Next.ttc", 2),
            ("/System/Library/Fonts/Supplemental/Arial Bold.ttf", None),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", None),
        ],
        "BookSans-Italic": [
            ("/System/Library/Fonts/Avenir Next.ttc", 4),
            ("/System/Library/Fonts/Supplemental/Arial Italic.ttf", None),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf", None),
        ],
        "BookMono": [
            ("/System/Library/Fonts/Menlo.ttc", 0),
            ("/System/Library/Fonts/Supplemental/Courier New.ttf", None),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf", None),
        ],
        "BookMono-Bold": [
            ("/System/Library/Fonts/Menlo.ttc", 1),
            ("/System/Library/Fonts/Supplemental/Courier New Bold.ttf", None),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf", None),
        ],
    }
    resolved: dict[str, str] = {}
    for logical, options in candidates.items():
        for raw, subfont_index in options:
            path = Path(raw)
            if not path.exists():
                continue
            try:
                kwargs = {"subfontIndex": subfont_index} if subfont_index is not None else {}
                pdfmetrics.registerFont(TTFont(logical, str(path), **kwargs))
            except Exception:
                continue
            resolved[logical] = logical
            break

    body = resolved.get("BookBody", resolved.get("BookSans", "Times-Roman"))
    body_bold = resolved.get("BookBody-Bold", resolved.get("BookSans-Bold", "Times-Bold"))
    body_italic = resolved.get("BookBody-Italic", resolved.get("BookSans-Italic", "Times-Italic"))
    sans = resolved.get("BookSans", "Helvetica")
    bold = resolved.get("BookSans-Bold", "Helvetica-Bold")
    italic = resolved.get("BookSans-Italic", "Helvetica-Oblique")
    mono = resolved.get("BookMono", "Courier")
    mono_bold = resolved.get("BookMono-Bold", "Courier-Bold")
    pdfmetrics.registerFontFamily("BookBody", normal=body, bold=body_bold, italic=body_italic, boldItalic=body_bold)
    pdfmetrics.registerFontFamily("BookSans", normal=sans, bold=bold, italic=italic, boldItalic=bold)
    pdfmetrics.registerFontFamily("BookMono", normal=mono, bold=mono_bold, italic=mono, boldItalic=mono_bold)
    return {
        "body": body,
        "body_bold": body_bold,
        "body_italic": body_italic,
        "sans": sans,
        "bold": bold,
        "italic": italic,
        "mono": mono,
        "mono_bold": mono_bold,
    }


def make_pdf_styles(fonts: dict[str, str]) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "body": ParagraphStyle("BookBody", parent=base["BodyText"], fontName=fonts["body"], fontSize=10.4, leading=14.6, textColor=INK, spaceAfter=8, alignment=TA_LEFT, allowWidows=0, allowOrphans=0),
        "small": ParagraphStyle("BookSmall", parent=base["BodyText"], fontName=fonts["sans"], fontSize=8.6, leading=11.3, textColor=MUTED, spaceAfter=5),
        "cover_kicker": ParagraphStyle("CoverKicker", parent=base["BodyText"], fontName=fonts["bold"], fontSize=10, leading=12, textColor=GOLD_TEXT, alignment=TA_CENTER, spaceAfter=18),
        "cover_title": ParagraphStyle("CoverTitle", parent=base["Title"], fontName=fonts["bold"], fontSize=31, leading=36, textColor=NAVY, alignment=TA_CENTER, spaceAfter=12),
        "cover_subtitle": ParagraphStyle("CoverSubtitle", parent=base["BodyText"], fontName=fonts["sans"], fontSize=15, leading=20, textColor=DARK_BLUE, alignment=TA_CENTER, spaceAfter=20),
        "cover_meta": ParagraphStyle("CoverMeta", parent=base["BodyText"], fontName=fonts["sans"], fontSize=9.5, leading=13, textColor=MUTED, alignment=TA_CENTER, spaceAfter=8),
        "cover_banner": ParagraphStyle("CoverBanner", parent=base["BodyText"], fontName=fonts["bold"], fontSize=9.5, leading=12, textColor=colors.white, alignment=TA_CENTER, spaceAfter=0),
        "cover_author_label": ParagraphStyle("CoverAuthorLabel", parent=base["BodyText"], fontName=fonts["bold"], fontSize=8.5, leading=10, textColor=GOLD_TEXT, alignment=TA_CENTER, spaceAfter=3),
        "cover_author": ParagraphStyle("CoverAuthor", parent=base["BodyText"], fontName=fonts["bold"], fontSize=14, leading=17, textColor=NAVY, alignment=TA_CENTER, spaceAfter=15),
        "cover_deck": ParagraphStyle("CoverDeck", parent=base["BodyText"], fontName=fonts["sans"], fontSize=9.8, leading=13, textColor=MUTED, alignment=TA_CENTER, leftIndent=24, rightIndent=24, spaceAfter=0),
        "cover_scope": ParagraphStyle("CoverScope", parent=base["BodyText"], fontName=fonts["bold"], fontSize=8.5, leading=11, textColor=NAVY, alignment=TA_CENTER, leftIndent=(CONTENT_W - COVER_W) / 2, rightIndent=(CONTENT_W - COVER_W) / 2, spaceAfter=0),
        "cover_stat_line": ParagraphStyle("CoverStatLine", parent=base["BodyText"], fontName=fonts["bold"], fontSize=8.2, leading=10, textColor=MUTED, alignment=TA_CENTER, leftIndent=(CONTENT_W - COVER_W) / 2, rightIndent=(CONTENT_W - COVER_W) / 2, spaceAfter=0),
        "part": ParagraphStyle("PartTitle", parent=base["Title"], fontName=fonts["bold"], fontSize=25, leading=31, textColor=NAVY, alignment=TA_CENTER, spaceAfter=20),
        "h1": ParagraphStyle("ChapterTitle", parent=base["Heading1"], fontName=fonts["bold"], fontSize=20.5, leading=25, textColor=NAVY, spaceBefore=0, spaceAfter=13, keepWithNext=1),
        "h2": ParagraphStyle("SectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.4, leading=17.2, textColor=BLUE, spaceBefore=14, spaceAfter=7, keepWithNext=1),
        "h3": ParagraphStyle("SubsectionTitle", parent=base["Heading3"], fontName=fonts["bold"], fontSize=11.2, leading=14.6, textColor=DARK_BLUE, spaceBefore=10, spaceAfter=5, keepWithNext=1),
        "h4": ParagraphStyle("MinorTitle", parent=base["Heading4"], fontName=fonts["bold"], fontSize=10.3, leading=13.4, textColor=DARK_BLUE, spaceBefore=8, spaceAfter=4, keepWithNext=1),
        "pattern_h2": ParagraphStyle("PatternSectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.2, leading=16.6, textColor=NAVY, backColor=PATTERN_BG, borderColor=CYAN, borderWidth=0.6, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "decision_h2": ParagraphStyle("DecisionSectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.2, leading=16.6, textColor=NAVY, backColor=PALE, borderColor=BLUE, borderWidth=0.6, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "checkpoint_h2": ParagraphStyle("CheckpointSectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.0, leading=16.4, textColor=NAVY, backColor=GUIDE_BG, borderColor=GOLD, borderWidth=0.6, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "practice_h2": ParagraphStyle("PracticeSectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.0, leading=16.4, textColor=NAVY, backColor=PRACTICE_BG, borderColor=GREEN, borderWidth=0.6, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "warning_h2": ParagraphStyle("WarningSectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.0, leading=16.4, textColor=NAVY, backColor=WARNING_BG, borderColor=RED, borderWidth=0.6, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "summary_h2": ParagraphStyle("SummarySectionTitle", parent=base["Heading2"], fontName=fonts["bold"], fontSize=13.0, leading=16.4, textColor=NAVY, backColor=SUMMARY_BG, borderColor=MUTED, borderWidth=0.5, borderRadius=4, borderPadding=(6, 8, 6, 8), spaceBefore=15, spaceAfter=9, keepWithNext=1),
        "walkthrough_h3": ParagraphStyle("WalkthroughTitle", parent=base["Heading3"], fontName=fonts["bold"], fontSize=10.5, leading=13.4, textColor=CYAN, spaceBefore=11, spaceAfter=5, keepWithNext=1),
        "blockquote": ParagraphStyle("Callout", parent=base["BodyText"], fontName=fonts["body"], fontSize=10.0, leading=14.2, textColor=DARK_BLUE, leftIndent=8, rightIndent=4, spaceAfter=0),
        "definition_body": ParagraphStyle("DefinitionBody", parent=base["BodyText"], fontName=fonts["body"], fontSize=10.4, leading=14.6, textColor=INK, leftIndent=14, spaceAfter=8, allowWidows=0, allowOrphans=0),
        "caption": ParagraphStyle("FigureCaption", parent=base["BodyText"], fontName=fonts["italic"], fontSize=8.5, leading=11, textColor=MUTED, alignment=TA_CENTER, spaceBefore=4, spaceAfter=10),
        "table_caption": ParagraphStyle("TableCaption", parent=base["BodyText"], fontName=fonts["bold"], fontSize=8.6, leading=11.2, textColor=DARK_BLUE, spaceAfter=5),
        "code": ParagraphStyle("Code", parent=base["Code"], fontName=fonts["mono"], fontSize=8.05, leading=10.35, textColor=INK, leftIndent=0, rightIndent=0, spaceAfter=0),
        "code_label": ParagraphStyle("CodeLabel", parent=base["BodyText"], fontName=fonts["bold"], fontSize=7.4, leading=9.2, textColor=colors.white, spaceAfter=0),
        "table": ParagraphStyle("TableText", parent=base["BodyText"], fontName=fonts["sans"], fontSize=8.7, leading=11.25, textColor=INK, spaceAfter=1),
        "table_head": ParagraphStyle("TableHead", parent=base["BodyText"], fontName=fonts["bold"], fontSize=8.7, leading=11.25, textColor=colors.white, spaceAfter=1),
        "toc_title": ParagraphStyle("TOCTitle", parent=base["Title"], fontName=fonts["bold"], fontSize=24, leading=29, textColor=NAVY, spaceAfter=18),
        "copyright": ParagraphStyle("Copyright", parent=base["BodyText"], fontName=fonts["sans"], fontSize=9.5, leading=13, textColor=INK, spaceAfter=8),
    }


class BookDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, styles: dict[str, ParagraphStyle], **kwargs: Any):
        super().__init__(filename, pagesize=letter, leftMargin=MARGIN, rightMargin=MARGIN, topMargin=62, bottomMargin=56, title=TITLE, author=AUTHOR, subject=SUBTITLE, **kwargs)
        self.styles_map = styles
        self.current_part = ""
        self.current_chapter = ""
        self.current_section = ""
        self.page_context: dict[int, str] = {}
        self._context_page = 0
        self._page_has_body = False
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body", leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
        self.addPageTemplates(
            PageTemplate(
                id="book",
                frames=[frame],
                onPage=self.capture_page_start,
                onPageEnd=self.draw_page,
            )
        )

    def beforeDocument(self) -> None:
        # multiBuild performs more than one pass; running-header state must not
        # leak from the end of an earlier pass into the front matter.
        self.current_part = ""
        self.current_chapter = ""
        self.current_section = ""
        self.page_context = {}
        self._context_page = 0
        self._page_has_body = False

    def capture_page_start(self, canvas, doc) -> None:
        """Freeze inherited context before any later heading on the page."""
        self._context_page = doc.page
        self._page_has_body = False
        inherited = self.current_section or self.current_chapter or self.current_part
        if inherited:
            self.page_context.setdefault(doc.page, inherited)

    def draw_page(self, canvas, doc) -> None:
        canvas.saveState()
        canvas.setTitle(TITLE)
        canvas.setAuthor(AUTHOR)
        if doc.page == 1:
            canvas.setFillColor(NAVY)
            canvas.rect(0, PAGE_H - 12, PAGE_W, 12, stroke=0, fill=1)
            canvas.setFillColor(GOLD)
            canvas.rect(0, 0, PAGE_W, 7, stroke=0, fill=1)
        else:
            canvas.setStrokeColor(LINE)
            canvas.setLineWidth(0.5)
            canvas.line(MARGIN, PAGE_H - 38, PAGE_W - MARGIN, PAGE_H - 38)
            canvas.setFont(self.styles_map["small"].fontName, 7.6)
            canvas.setFillColor(MUTED)
            canvas.drawString(MARGIN, PAGE_H - 31, TITLE)
            right = self.page_context.get(doc.page) or self.current_section or self.current_chapter or self.current_part or "SDE-2 Interview Preparation Guide"
            canvas.drawRightString(PAGE_W - MARGIN, PAGE_H - 31, right[:72])
            canvas.line(MARGIN, 39, PAGE_W - MARGIN, 39)
            canvas.drawRightString(PAGE_W - MARGIN, 27, str(doc.page))
            canvas.drawString(MARGIN, 27, "Java 21 baseline")
        canvas.restoreState()

    def afterFlowable(self, flowable: Flowable) -> None:
        if self._context_page != self.page:
            self._context_page = self.page
            self._page_has_body = False
        if not isinstance(flowable, Paragraph):
            if not isinstance(flowable, (ActionFlowable, PageBreak, CondPageBreak, Spacer, HRFlowable)):
                self._page_has_body = True
            return
        text = getattr(flowable, "_bookmark_text", flowable.getPlainText())
        heading_level = getattr(flowable, "_heading_level", None)
        if text.startswith("Part ") or text == "Appendices":
            self.current_part = text
            self.current_chapter = ""
            self.current_section = ""
            self.page_context[self.page] = text
        elif text.startswith("Chapter ") or text.startswith("Appendix "):
            self.current_chapter = text
            self.current_section = ""
            self.page_context[self.page] = text
        elif heading_level == 1:
            self.current_section = text
            if not self._page_has_body:
                self.page_context[self.page] = text
            else:
                self.page_context.setdefault(self.page, text)
        elif heading_level == 2:
            self.current_section = text
            if not self._page_has_body or getattr(flowable, "_force_page_context", False):
                self.page_context[self.page] = text
        if text.strip():
            self._page_has_body = True
        if not hasattr(flowable, "_bookmark_name"):
            return
        key = getattr(flowable, "_bookmark_name")
        toc_level = getattr(flowable, "_toc_level", 0)
        outline_level = getattr(flowable, "_outline_level", toc_level)
        self.canv.bookmarkPage(key)
        try:
            self.canv.addOutlineEntry(text, key, level=outline_level, closed=outline_level > 0)
        except ValueError:
            # A defensive fallback for an unexpected source heading jump.
            self.canv.addOutlineEntry(text, key, level=max(0, outline_level - 1), closed=True)
        if getattr(flowable, "_include_toc", True):
            self.notify("TOCEntry", (toc_level, text, self.page, key))


class GroupedTableOfContents(TableOfContents):
    """Repeat TOC context and keep each part with its first child entry."""

    def __init__(self, header_style: ParagraphStyle, **kwargs: Any):
        super().__init__(**kwargs)
        self.header_style = header_style

    def wrap(self, availWidth: float, availHeight: float) -> tuple[float, float]:
        super().wrap(availWidth, availHeight)
        entries = self._lastEntries or [(0, "Placeholder for table of contents", 0, None)]
        original_rows = list(self._table._cellvalues)
        rows = [[Paragraph("CONTENTS NAVIGATION", self.header_style)], *original_rows]
        commands: list[tuple[Any, ...]] = [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 1), (-1, -1), 0),
            ("RIGHTPADDING", (0, 1), (-1, -1), 0),
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("LEFTPADDING", (0, 0), (-1, 0), 7),
            ("RIGHTPADDING", (0, 0), (-1, 0), 7),
            ("TOPPADDING", (0, 0), (-1, 0), 4),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ]
        entry_rows: list[int] = []
        cursor = 1
        for level, _, _, _ in entries:
            style = self.getLevelStyle(level)
            if style.spaceBefore:
                cursor += 1
            entry_rows.append(cursor)
            cursor += 1
        for index, entry in enumerate(entries[:-1]):
            if entry[0] == 0 and entries[index + 1][0] > entry[0]:
                commands.append(
                    ("NOSPLIT", (0, entry_rows[index]), (0, entry_rows[index + 1]))
                )
        self._table = Table(
            rows,
            colWidths=(availWidth,),
            repeatRows=1,
            style=TableStyle(commands),
        )
        self.width, self.height = self._table.wrapOn(
            self.canv, availWidth, availHeight
        )
        return self.width, self.height


def bookmark_name(text: str, index: int) -> str:
    digest = hashlib.sha1(f"{index}:{text}".encode("utf-8")).hexdigest()[:12]
    return f"section-{digest}"


def plain_inlines(inlines: Sequence[dict[str, Any]]) -> str:
    parts: list[str] = []
    for item in inlines:
        kind = item.get("t")
        value = item.get("c")
        if kind == "Str":
            parts.append(str(value))
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            parts.append(" ")
        elif kind in ("Emph", "Strong", "Strikeout", "Superscript", "Subscript", "SmallCaps"):
            parts.append(plain_inlines(value))
        elif kind == "Code":
            parts.append(value[1])
        elif kind == "Link":
            parts.append(plain_inlines(value[1]))
        elif kind == "Image":
            parts.append(plain_inlines(value[1]))
        elif kind == "Quoted":
            parts.append('"' + plain_inlines(value[1]) + '"')
        elif kind == "Math":
            parts.append(value[1])
        elif kind == "Span":
            parts.append(plain_inlines(value[1]))
        elif kind == "RawInline":
            parts.append(re.sub(r"<[^>]+>", "", value[1]))
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def markup_inlines(inlines: Sequence[dict[str, Any]], fonts: dict[str, str]) -> str:
    parts: list[str] = []
    for item in inlines:
        kind = item.get("t")
        value = item.get("c")
        if kind == "Str":
            parts.append(html.escape(str(value)))
        elif kind in ("Space", "SoftBreak"):
            parts.append(" ")
        elif kind == "LineBreak":
            parts.append("<br/>")
        elif kind == "Emph":
            parts.append(f"<i>{markup_inlines(value, fonts)}</i>")
        elif kind == "Strong":
            parts.append(f"<b>{markup_inlines(value, fonts)}</b>")
        elif kind == "Strikeout":
            parts.append(f"<strike>{markup_inlines(value, fonts)}</strike>")
        elif kind == "Superscript":
            parts.append(f"<super>{markup_inlines(value, fonts)}</super>")
        elif kind == "Subscript":
            parts.append(f"<sub>{markup_inlines(value, fonts)}</sub>")
        elif kind == "Code":
            parts.append(f'<font name="{fonts["mono"]}" color="#1F4D78">{html.escape(value[1])}</font>')
        elif kind == "Link":
            label = markup_inlines(value[1], fonts)
            url = html.escape(value[2][0], quote=True)
            parts.append(f'<link href="{url}" color="#1F5A94">{label}</link>')
        elif kind == "Image":
            parts.append(html.escape(plain_inlines(value[1])))
        elif kind == "Quoted":
            parts.append('&quot;' + markup_inlines(value[1], fonts) + '&quot;')
        elif kind == "Math":
            parts.append(html.escape(value[1]))
        elif kind == "Span":
            parts.append(markup_inlines(value[1], fonts))
        elif kind == "RawInline":
            parts.append(html.escape(re.sub(r"<[^>]+>", "", value[1])))
    return "".join(parts)


JAVA_KEYWORDS = {
    "abstract", "assert", "boolean", "break", "byte", "case", "catch", "char", "class", "const", "continue", "default", "do", "double", "else", "enum", "extends", "final", "finally", "float", "for", "goto", "if", "implements", "import", "instanceof", "int", "interface", "long", "native", "new", "non-sealed", "package", "permits", "private", "protected", "public", "record", "return", "sealed", "short", "static", "strictfp", "super", "switch", "synchronized", "this", "throw", "throws", "transient", "try", "var", "void", "volatile", "while", "yield", "true", "false", "null",
}
TOKEN_RE = re.compile(
    r"(?P<comment>//[^\n]*|/\*.*?\*/)|(?P<string>\"(?:\\.|[^\"\\])*\"|'(?:\\.|[^'\\])*')|(?P<annotation>@[A-Za-z_$][\w$]*)|(?P<number>\b(?:0[xX][0-9A-Fa-f_]+|0[bB][01_]+|\d[\d_]*(?:\.\d[\d_]*)?(?:[eE][+-]?\d+)?[fFdDlL]?)\b)|(?P<word>\b[A-Za-z_$][\w$]*(?:-[A-Za-z]+)?\b)",
    re.DOTALL | re.MULTILINE,
)


def highlighted_code(code: str, language: str) -> str:
    code = code.replace("\t", "    ")
    if language.lower() not in {"java", "jshell"}:
        return html.escape(code)
    output: list[str] = []
    pos = 0
    for match in TOKEN_RE.finditer(code):
        output.append(html.escape(code[pos : match.start()]))
        token = html.escape(match.group(0))
        kind = match.lastgroup
        if kind == "comment":
            output.append(f'<font color="#2D7D66">{token}</font>')
        elif kind == "string":
            output.append(f'<font color="#8A4B08">{token}</font>')
        elif kind == "annotation":
            output.append(f'<font color="#7A3E9D">{token}</font>')
        elif kind == "number":
            output.append(f'<font color="#7A3E9D">{token}</font>')
        elif kind == "word" and match.group(0) in JAVA_KEYWORDS:
            output.append(f'<font color="#1F5A94"><b>{token}</b></font>')
        else:
            output.append(token)
        pos = match.end()
    output.append(html.escape(code[pos:]))
    return "".join(output)


def extract_blocks_text(blocks: Sequence[dict[str, Any]]) -> str:
    pieces: list[str] = []
    for block in blocks:
        kind = block.get("t")
        value = block.get("c")
        if kind in ("Para", "Plain"):
            pieces.append(plain_inlines(value))
        elif kind == "Header":
            pieces.append(plain_inlines(value[2]))
        elif kind == "CodeBlock":
            pieces.append(value[1])
        elif kind in ("BulletList",):
            for item in value:
                pieces.append(extract_blocks_text(item))
        elif kind == "OrderedList":
            for item in value[1]:
                pieces.append(extract_blocks_text(item))
        elif kind == "BlockQuote":
            pieces.append(extract_blocks_text(value))
    return " ".join(pieces).strip()


def image_flowable(target: str, caption: str, styles: dict[str, ParagraphStyle]) -> list[Flowable]:
    target = target.split("#", 1)[0]
    path = Path(target)
    if not path.is_absolute():
        path = ROOT / path
    if not path.exists():
        return [Paragraph(f"[Missing figure: {html.escape(caption)}]", styles["blockquote"])]
    with PILImage.open(path) as image:
        width, height = image.size
    max_width = CONTENT_W - 8
    max_height = 330
    scale = min(max_width / width, max_height / height)
    rendered = RLImage(str(path), width=width * scale, height=height * scale)
    rendered.hAlign = "CENTER"
    items: list[Flowable] = [rendered]
    if caption:
        # Use one native keep-with-next chain instead of nesting a
        # KeepTogether inside a heading's own keep chain. ReportLab measures a
        # nested KeepTogether as an intentionally huge flowable, which can
        # leave the heading behind when the figure moves to the next page.
        rendered.keepWithNext = 1
        items.append(Paragraph(html.escape(caption), styles["caption"]))
    return items


def code_chunks(lines: list[str], maximum: int = 36) -> list[list[str]]:
    """Balance a listing across the fewest panels that stay under the cap."""
    chunk_count = max(1, math.ceil(len(lines) / maximum))
    base_size, larger_chunks = divmod(len(lines), chunk_count)
    sizes = [base_size + 1] * larger_chunks + [base_size] * (chunk_count - larger_chunks)
    chunks: list[list[str]] = []
    offset = 0
    for size in sizes:
        chunks.append(lines[offset : offset + size])
        offset += size
    return chunks


def code_flowables(code: str, language: str, styles: dict[str, ParagraphStyle]) -> list[Flowable]:
    lines = code.splitlines() or [""]
    chunks = code_chunks(lines)
    result: list[Flowable] = []
    for idx, chunk in enumerate(chunks):
        label = language.upper() if language else "CODE"
        if len(chunks) > 1:
            label += f" (continued {idx + 1}/{len(chunks)})"
        label_flowable = Paragraph(html.escape(label), styles["code_label"])
        markup = highlighted_code("\n".join(chunk), language)
        longest_line = max((len(line.expandtabs(4)) for line in chunk), default=0)
        code_style = styles["code"]
        if longest_line > 82:
            code_style = ParagraphStyle(
                "CodeLongLine",
                parent=styles["code"],
                fontSize=7.3,
                leading=9.5,
            )
        pre = XPreformatted(markup, code_style)
        panel = Table(
            [[label_flowable], [pre]],
            colWidths=[CONTENT_W],
            hAlign="LEFT",
            splitByRow=0,
            splitInRow=0,
        )
        panel.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                    ("BACKGROUND", (0, 1), (-1, -1), LIGHT),
                    ("BOX", (0, 0), (-1, -1), 0.5, LINE),
                    ("LINEBEFORE", (0, 1), (0, -1), 2.5, CYAN),
                    ("LEFTPADDING", (0, 0), (-1, 0), 8),
                    ("RIGHTPADDING", (0, 0), (-1, 0), 8),
                    ("TOPPADDING", (0, 0), (-1, 0), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
                    ("LEFTPADDING", (0, 1), (-1, -1), 9),
                    ("RIGHTPADDING", (0, 1), (-1, -1), 8),
                    ("TOPPADDING", (0, 1), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 1), (-1, -1), 7),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        # The two-row table is deliberately unsplittable, so its language
        # label and code remain one panel. Keeping it as a native Table also
        # lets a preceding heading or list prompt form a correct keep chain.
        result.extend([panel, Spacer(1, 9)])
    return result


def table_flowable(
    value: Any,
    styles: dict[str, ParagraphStyle],
    fonts: dict[str, str],
    *,
    keep_short: bool = True,
) -> list[Flowable]:
    attr, caption, colspecs, thead, tbodies, tfoot = value
    rows_raw: list[Any] = []
    head_rows = thead[1] if thead and len(thead) > 1 else []
    rows_raw.extend(head_rows)
    for body in tbodies:
        if len(body) >= 4:
            rows_raw.extend(body[2])
            rows_raw.extend(body[3])
    if tfoot and len(tfoot) > 1:
        rows_raw.extend(tfoot[1])
    if not rows_raw:
        return []

    col_count = len(colspecs)
    data: list[list[Any]] = []
    lengths = [3] * col_count
    for row_idx, row in enumerate(rows_raw):
        cells = row[1] if isinstance(row, list) and len(row) > 1 else []
        converted: list[Any] = []
        for col_idx in range(col_count):
            if col_idx < len(cells):
                cell = cells[col_idx]
                blocks = cell[4] if len(cell) > 4 else []
                text = extract_blocks_text(blocks)
                lengths[col_idx] = max(lengths[col_idx], min(100, len(text)))
                cell_style = styles["table_head"] if row_idx < len(head_rows) else styles["table"]
                flowables: list[Flowable] = []
                for block in blocks:
                    if block.get("t") in ("Plain", "Para"):
                        flowables.append(Paragraph(markup_inlines(block["c"], fonts), cell_style))
                    elif block.get("t") == "CodeBlock":
                        flowables.append(XPreformatted(html.escape(block["c"][1]), styles["code"]))
                converted.append(flowables or Paragraph("", cell_style))
            else:
                converted.append(Paragraph("", styles["table"]))
        data.append(converted)

    weights = [max(3.0, math.sqrt(length)) for length in lengths]
    total = sum(weights)
    widths = [CONTENT_W * weight / total for weight in weights]
    minimum = 44
    for idx, width in enumerate(widths):
        if width < minimum:
            deficit = minimum - width
            widths[idx] = minimum
            largest = max(range(len(widths)), key=widths.__getitem__)
            if largest != idx:
                widths[largest] -= deficit

    body_row_count = len(data) - len(head_rows)
    split_kwargs: dict[str, Any] = {}
    if body_row_count >= 9:
        split_kwargs["rowSplitRange"] = (len(head_rows) + 4, -4)
    table = LongTable(
        data,
        colWidths=widths,
        repeatRows=len(head_rows),
        hAlign="LEFT",
        splitByRow=1,
        splitInRow=0,
        **split_kwargs,
    )
    commands = [
        ("BOX", (0, 0), (-1, -1), 0.45, LINE),
        ("LINEBELOW", (0, 0), (-1, -1), 0.3, LINE),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    if head_rows:
        commands.append(("BACKGROUND", (0, 0), (-1, len(head_rows) - 1), NAVY))
        if len(data) > len(head_rows):
            commands.append(
                ("ROWBACKGROUNDS", (0, len(head_rows)), (-1, -1), [colors.white, TABLE_ALT])
            )
    table.setStyle(TableStyle(commands))
    long_caption = caption[1] if caption and len(caption) > 1 else []
    caption_text = extract_blocks_text(long_caption)
    table_content: list[Flowable] = []
    if caption_text:
        table_caption = Paragraph(html.escape(caption_text), styles["table_caption"])
        table_caption.keepWithNext = 1
        table_content.append(table_caption)
    else:
        table.spaceBefore = 4
    table_content.append(table)

    # Short tables are easier to understand when they stay on one page. Long
    # tables still split by whole rows, repeat their dark header, and reserve
    # enough space for a meaningful first segment before they begin.
    _, table_height = table.wrap(CONTENT_W, FRAME_H)
    if table_height <= FRAME_H * 0.72:
        # A short table should move as a whole. Avoid wrapping it in
        # KeepTogether: a heading followed by that nested wrapper can be
        # stranded on the previous page by ReportLab's sentinel height.
        table.splitByRow = 0
        table.splitInRow = 0
        return [*table_content, Spacer(1, 10)]
    return [CondPageBreak(FRAME_H * 0.32), *table_content, Spacer(1, 10)]


def task_list_flowable(
    raw_items: Sequence[Sequence[dict[str, Any]]],
    styles: dict[str, ParagraphStyle],
    fonts: dict[str, str],
) -> Flowable | None:
    """Render GFM task lists without a redundant bullet or extractor-hostile glyph."""
    rows: list[list[Flowable]] = []
    for item in raw_items:
        if len(item) != 1 or item[0].get("t") not in ("Plain", "Para"):
            return None
        inlines = list(item[0].get("c") or [])
        if not inlines or inlines[0].get("t") != "Str":
            return None
        marker = inlines[0].get("c")
        if marker not in ("☐", "☒"):
            return None
        inlines = inlines[1:]
        if inlines and inlines[0].get("t") == "Space":
            inlines = inlines[1:]
        ascii_marker = "[ ]" if marker == "☐" else "[x]"
        rows.append(
            [Paragraph(f"<b>{ascii_marker}</b>&nbsp;&nbsp;{markup_inlines(inlines, fonts)}", styles["body"])]
        )
    if not rows:
        return None
    table = LongTable(rows, colWidths=[CONTENT_W - 21], hAlign="LEFT", splitByRow=1)
    table.setStyle(
        TableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 13),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return table


class PdfRenderer:
    def __init__(self, styles: dict[str, ParagraphStyle], fonts: dict[str, str]):
        self.styles = styles
        self.fonts = fonts
        self.heading_index = 0
        self.in_part = False

    def heading(self, level: int, text: str) -> list[Flowable]:
        self.heading_index += 1
        name = bookmark_name(text, self.heading_index)
        if level == 1 and (text.startswith("Part ") or text == "Appendices"):
            self.in_part = True
            para = Paragraph(html.escape(text), self.styles["part"])
            para._bookmark_name = name
            para._bookmark_text = text
            para._toc_level = 0
            para._outline_level = 0
            para._heading_level = 1
            # The following chapter/appendix H1 supplies the page break after
            # this divider.  Ending here with another PageBreak would combine
            # with that H1's leading break and create a blank page.
            return [CondPageBreak(FRAME_H - 12), Spacer(1, 145), Paragraph("JAVA FOUNDATIONS TO ADVANCED ENGINEERING", self.styles["cover_kicker"]), para, HRFlowable(width="45%", thickness=2, color=GOLD, spaceBefore=8, spaceAfter=12), Paragraph("A focused part of the complete SDE-2 preparation guide", self.styles["cover_meta"])]

        if level == 1:
            is_chapter = text.startswith("Chapter ")
            is_appendix = text.startswith("Appendix ")
            toc_level = 1 if (is_chapter or is_appendix) else 0
            outline_level = toc_level
            para = Paragraph(html.escape(text), self.styles["h1"])
            para._bookmark_name = name
            para._bookmark_text = text
            para._toc_level = toc_level
            para._outline_level = outline_level
            para._heading_level = 1
            # Start on a fresh page unless the preceding flowable already
            # filled its frame and advanced us to the top of a new page. A
            # hard PageBreak in that state creates a running-header-only page.
            # Allow the invisible 4–10pt trailing spacer from the preceding
            # panel to be present at frame top without manufacturing a blank
            # page; any real line of content still forces a fresh chapter page.
            items: list[Flowable] = [CondPageBreak(FRAME_H - 12)]
            if is_chapter:
                items.append(Paragraph("SDE-2 CORE CHAPTER", self.styles["cover_kicker"]))
            elif is_appendix:
                items.append(Paragraph("REFERENCE APPENDIX", self.styles["cover_kicker"]))
            items.extend([para, HRFlowable(width="100%", thickness=1.2, color=GOLD, spaceBefore=0, spaceAfter=12)])
            return items

        lowered = text.casefold()
        semantic = re.sub(r"^\d+(?:\.\d+)*\s*", "", lowered).strip()
        display_text = text
        if level == 2 and re.match(r"^(?:pattern|family)\s+\d+\b", semantic):
            style = self.styles["pattern_h2"]
        elif level == 2 and any(
            signal in semantic
            for signal in ("decision map", "decision guide", "selection guide", "recognition map")
        ):
            style = self.styles["decision_h2"]
            display_text = f"CHOOSE IT | {text}"
        elif level == 2 and any(
            signal in semantic
            for signal in (
                "first-principles",
                "first principles",
                "invariant",
                "proof",
                "why this matters",
            )
        ):
            style = self.styles["decision_h2"]
            display_text = f"WHY IT WORKS | {text}"
        elif level == 2 and any(
            signal in semantic
            for signal in (
                "worked java example",
                "worked example",
                "execution or memory walkthrough",
                "execution walkthrough",
                "worked trace",
                "dry run",
            )
        ):
            style = self.styles["pattern_h2"]
            display_text = f"TRACE IT | {text}"
        elif level == 2 and any(
            signal in semantic
            for signal in (
                "edge cases",
                "common mistakes",
                "failure modes",
                "pitfalls",
            )
        ):
            style = self.styles["warning_h2"]
            display_text = f"WATCH OUT | {text}"
        elif level == 2 and any(
            signal in semantic
            for signal in (
                "completion check",
                "practice ladder",
                "interview questions",
                "revision checklist",
                "exercises",
                "checkpoint",
            )
        ):
            style = self.styles["practice_h2"]
            display_text = f"TRY IT | {text}"
        elif level == 2 and any(
            signal in semantic for signal in ("chapter summary", "summary", "recap")
        ):
            style = self.styles["summary_h2"]
            display_text = f"RECAP | {text}"
        elif level == 3 and any(
            semantic.startswith(signal)
            for signal in ("dry run", "worked trace", "walkthrough", "execution trace")
        ):
            style = self.styles["walkthrough_h3"]
            display_text = f"TRACE IT | {text}"
        else:
            style = self.styles["h2"] if level == 2 else self.styles["h3"] if level == 3 else self.styles["h4"]
        para = Paragraph(html.escape(display_text), style)
        para._bookmark_name = name
        para._bookmark_text = text
        para._heading_level = level
        para._include_toc = False
        para._toc_level = 2
        para._outline_level = 2 if self.in_part else 1
        minimum_follow_space = 145 if level == 2 else 100 if level == 3 else 78
        return [CondPageBreak(minimum_follow_space), para]

    def list_item(self, blocks: Sequence[dict[str, Any]]) -> list[Flowable]:
        items: list[Flowable] = []
        for block in blocks:
            kind = block.get("t")
            if kind in ("Para", "Plain"):
                items.append(Paragraph(markup_inlines(block["c"], self.fonts), self.styles["body"]))
            elif kind == "CodeBlock":
                items.extend(code_flowables(block["c"][1], (block["c"][0][1] or [""])[0], self.styles))
            elif kind == "BulletList":
                task_list = task_list_flowable(block["c"], self.styles, self.fonts)
                if task_list is not None:
                    items.append(task_list)
                else:
                    children = [ListItem(self.list_item(child), leftIndent=12) for child in block["c"]]
                    items.append(ListFlowable(children, bulletType="bullet", leftIndent=18, bulletFontName=self.fonts["sans"], bulletFontSize=7))
            elif kind == "OrderedList":
                children = [ListItem(self.list_item(child), leftIndent=12) for child in block["c"][1]]
                items.append(ListFlowable(children, bulletType="1", start=block["c"][0][0], leftIndent=20, bulletFontName=self.fonts["sans"], bulletFontSize=8))
        return items or [Paragraph("", self.styles["body"])]

    def blocks(self, blocks: Sequence[dict[str, Any]]) -> list[Flowable]:
        story: list[Flowable] = []
        skip_until = -1
        for block_index, block in enumerate(blocks):
            if block_index < skip_until:
                continue
            kind = block.get("t")
            value = block.get("c")
            if kind == "Header" and value[0] in (2, 3, 4):
                # Consecutive subheadings are one navigation chain. Only the
                # first needs a conditional page break; another CondPageBreak
                # between H2 and H3 would interrupt keepWithNext and could
                # strand the H2 above the page break.
                header_index = block_index
                heading_flowables: list[Flowable] = []
                first_heading = True
                while (
                    header_index < len(blocks)
                    and blocks[header_index].get("t") == "Header"
                    and blocks[header_index]["c"][0] in (2, 3, 4)
                ):
                    header = blocks[header_index]["c"]
                    rendered_heading = self.heading(
                        header[0], plain_inlines(header[2])
                    )
                    if not first_heading:
                        rendered_heading = [
                            item
                            for item in rendered_heading
                            if not isinstance(item, CondPageBreak)
                        ]
                    heading_flowables.extend(rendered_heading)
                    first_heading = False
                    header_index += 1

                # If a short lead-in introduces an atomic visual unit, keep
                # the complete opening together: heading chain, lead-in, and
                # the first figure/code/table/callout/list panel. Ordinary
                # prose still flows naturally after its first paragraph.
                payload_index = header_index
                setup_blocks: list[dict[str, Any]] = []
                while (
                    payload_index < len(blocks)
                    and len(setup_blocks) < 2
                    and blocks[payload_index].get("t") in ("Para", "Plain")
                    and not (
                        len(blocks[payload_index].get("c") or []) == 1
                        and blocks[payload_index]["c"][0].get("t") == "Image"
                    )
                ):
                    setup_blocks.append(blocks[payload_index])
                    payload_index += 1

                protected_kinds = {
                    "Table",
                    "CodeBlock",
                    "BlockQuote",
                    "BulletList",
                    "OrderedList",
                    "Figure",
                }
                payload_is_image = (
                    payload_index < len(blocks)
                    and blocks[payload_index].get("t") in ("Para", "Plain")
                    and len(blocks[payload_index].get("c") or []) == 1
                    and blocks[payload_index]["c"][0].get("t") == "Image"
                )
                payload_is_protected = (
                    payload_index < len(blocks)
                    and (
                        blocks[payload_index].get("t") in protected_kinds
                        or payload_is_image
                    )
                )
                if payload_is_protected:
                    setup_flowables = self.blocks(setup_blocks)
                    payload_end = payload_index + 1
                    payload_flowables = self.blocks([blocks[payload_index]])

                    # A one-item prompt followed by code is a single reading
                    # unit even when Pandoc emits it as two top-level blocks.
                    payload_kind = blocks[payload_index].get("t")
                    payload_value = blocks[payload_index].get("c")
                    single_item_list = (
                        payload_kind == "BulletList" and len(payload_value) == 1
                    ) or (
                        payload_kind == "OrderedList" and len(payload_value[1]) == 1
                    )
                    if (
                        single_item_list
                        and payload_end < len(blocks)
                        and blocks[payload_end].get("t") == "CodeBlock"
                    ):
                        payload_flowables = [
                            item for item in payload_flowables if not isinstance(item, Spacer)
                        ]
                        code_attr, code = blocks[payload_end]["c"]
                        language = (code_attr[1] or [""])[0]
                        payload_flowables.extend(
                            code_flowables(code, language, self.styles)
                        )
                        payload_end += 1

                    payload_flowables = [
                        item
                        for item in payload_flowables
                        if not isinstance(item, CondPageBreak)
                    ]

                    first_spacing = next(
                        (
                            index
                            for index, item in enumerate(payload_flowables)
                            if isinstance(item, Spacer)
                        ),
                        len(payload_flowables),
                    )
                    attached_payload = payload_flowables[:first_spacing]
                    remainder = payload_flowables[first_spacing:]
                    page_guards = [
                        item
                        for item in heading_flowables
                        if isinstance(item, CondPageBreak)
                    ]
                    heading_content = [
                        item
                        for item in heading_flowables
                        if not isinstance(item, CondPageBreak)
                    ]
                    story.extend(page_guards)
                    story.append(
                        KeepTogether(
                            [*heading_content, *setup_flowables, *attached_payload]
                        )
                    )
                    story.extend(remainder)
                    skip_until = payload_end
                    continue
                story.extend(heading_flowables)
                skip_until = header_index
                continue
            if kind == "Header":
                story.extend(self.heading(value[0], plain_inlines(value[2])))
            elif kind in ("Para", "Plain"):
                if len(value) == 1 and value[0].get("t") == "Image":
                    image = value[0]["c"]
                    caption = plain_inlines(image[1])
                    story.extend(image_flowable(image[2][0], caption, self.styles))
                else:
                    markup = markup_inlines(value, self.fonts)
                    if markup.strip():
                        # Let ordinary prose use the remaining frame naturally.
                        # Paragraphs already split only at line boundaries; wrapping
                        # every short paragraph in KeepTogether caused large holes
                        # before protected tables and code examples.
                        story.append(Paragraph(markup, self.styles["body"]))
            elif kind == "CodeBlock":
                attr, code = value
                language = (attr[1] or [""])[0]
                story.extend(code_flowables(code, language, self.styles))
            elif kind == "BulletList":
                task_list = task_list_flowable(value, self.styles, self.fonts)
                if task_list is not None:
                    paired_code = (
                        len(value) == 1
                        and block_index + 1 < len(blocks)
                        and blocks[block_index + 1].get("t") == "CodeBlock"
                    )
                    if paired_code:
                        code_attr, code = blocks[block_index + 1]["c"]
                        language = (code_attr[1] or [""])[0]
                        rendered_code = code_flowables(code, language, self.styles)
                        story.append(KeepTogether([task_list, rendered_code[0]]))
                        story.extend(rendered_code[1:])
                        skip_until = block_index + 2
                    else:
                        story.extend([task_list, Spacer(1, 4)])
                else:
                    items = [ListItem(self.list_item(item), leftIndent=13) for item in value]
                    list_flowable = ListFlowable(items, bulletType="bullet", leftIndent=21, bulletFontName=self.fonts["sans"], bulletFontSize=7, spaceAfter=6)
                    paired_code = (
                        len(value) == 1
                        and block_index + 1 < len(blocks)
                        and blocks[block_index + 1].get("t") == "CodeBlock"
                    )
                    if paired_code:
                        code_attr, code = blocks[block_index + 1]["c"]
                        language = (code_attr[1] or [""])[0]
                        rendered_code = code_flowables(code, language, self.styles)
                        story.append(KeepTogether([list_flowable, rendered_code[0]]))
                        story.extend(rendered_code[1:])
                        skip_until = block_index + 2
                    else:
                        story.append(list_flowable)
            elif kind == "OrderedList":
                attrs, raw_items = value
                items = [ListItem(self.list_item(item), leftIndent=13) for item in raw_items]
                list_flowable = ListFlowable(items, bulletType="1", start=attrs[0], leftIndent=23, bulletFontName=self.fonts["sans"], bulletFontSize=8, spaceAfter=6)
                paired_code = (
                    len(raw_items) == 1
                    and block_index + 1 < len(blocks)
                    and blocks[block_index + 1].get("t") == "CodeBlock"
                )
                if paired_code:
                    code_attr, code = blocks[block_index + 1]["c"]
                    language = (code_attr[1] or [""])[0]
                    rendered_code = code_flowables(code, language, self.styles)
                    story.append(KeepTogether([list_flowable, rendered_code[0]]))
                    story.extend(rendered_code[1:])
                    skip_until = block_index + 2
                else:
                    story.append(list_flowable)
            elif kind == "BlockQuote":
                quote_flowables: list[Flowable] = []
                for quote_block in value:
                    if quote_block.get("t") in ("Para", "Plain"):
                        quote_flowables.append(
                            Paragraph(
                                markup_inlines(quote_block["c"], self.fonts),
                                self.styles["blockquote"],
                            )
                        )
                    else:
                        quote_flowables.extend(self.blocks([quote_block]))
                if not quote_flowables:
                    quote_flowables.append(Paragraph("", self.styles["blockquote"]))
                callout = Table([[quote_flowables]], colWidths=[CONTENT_W - 10], hAlign="LEFT")
                callout.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), PATTERN_BG), ("LINEBEFORE", (0, 0), (0, -1), 3, CYAN), ("BOX", (0, 0), (-1, -1), 0.35, LINE), ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
                story.extend([callout, Spacer(1, 7)])
            elif kind == "HorizontalRule":
                story.append(HRFlowable(width="100%", thickness=0.5, color=LINE, spaceBefore=5, spaceAfter=8))
            elif kind == "Table":
                story.extend(table_flowable(value, self.styles, self.fonts))
            elif kind == "Div":
                story.extend(self.blocks(value[1]))
            elif kind == "DefinitionList":
                for term, defs in value:
                    story.append(Paragraph(f"<b>{markup_inlines(term, self.fonts)}</b>", self.styles["body"]))
                    for definition in defs:
                        for definition_block in definition:
                            if definition_block.get("t") in ("Para", "Plain"):
                                story.append(
                                    Paragraph(
                                        markup_inlines(definition_block["c"], self.fonts),
                                        self.styles["definition_body"],
                                    )
                                )
                            else:
                                story.extend(self.blocks([definition_block]))
            elif kind == "Figure":
                # Pandoc 3 figure: render the body; image paragraphs carry the caption.
                story.extend(self.blocks(value[2]))
            elif kind == "RawBlock":
                raw = value[1] if isinstance(value, list) and len(value) > 1 else ""
                if "PAGE_BREAK" in raw:
                    story.append(PageBreak())
            elif kind == "LineBlock":
                for line in value:
                    story.append(Paragraph(markup_inlines(line, self.fonts), self.styles["body"]))
        return story


class CoverBand(Flowable):
    def __init__(self, text: str, font_name: str):
        super().__init__()
        self.text = text
        self.font_name = font_name
        self.width = COVER_W
        self.height = 30
        self.hAlign = "CENTER"

    def wrap(self, avail_width: float, avail_height: float) -> tuple[float, float]:
        return min(self.width, avail_width), self.height

    def draw(self) -> None:
        self.canv.setFillColor(NAVY)
        self.canv.rect(0, 0, self.width, self.height, stroke=0, fill=1)
        self.canv.setFillColor(colors.white)
        self.canv.setFont(self.font_name, 9.5)
        self.canv.drawCentredString(self.width / 2, 10.5, self.text)


def cover_story(styles: dict[str, ParagraphStyle], fonts: dict[str, str]) -> list[Flowable]:
    return [
        Spacer(1, 10),
        CoverBand("JAVA | JVM | CONCURRENCY | SYSTEM DESIGN", fonts["bold"]),
        Spacer(1, 44),
        Paragraph("JAVA ENGINEERING INTERVIEW FIELD GUIDE", styles["cover_kicker"]),
        Paragraph(TITLE, styles["cover_title"]),
        Paragraph(SUBTITLE, styles["cover_subtitle"]),
        HRFlowable(width="48%", thickness=2, color=GOLD, spaceBefore=5, spaceAfter=16),
        Paragraph("BY", styles["cover_author_label"]),
        Paragraph(AUTHOR, styles["cover_author"]),
        Paragraph("A first-principles guide to Java semantics, JVM internals, collections, concurrency, performance, algorithms, and production engineering.", styles["cover_deck"]),
        Spacer(1, 21),
        HRFlowable(width=COVER_W, thickness=0.8, color=GOLD, spaceBefore=0, spaceAfter=7),
        Paragraph("LANGUAGE | JVM | COLLECTIONS | CONCURRENCY | DSA | BACKEND", styles["cover_scope"]),
        HRFlowable(width=COVER_W, thickness=0.55, color=LINE, spaceBefore=7, spaceAfter=11),
        Paragraph("54 CHAPTERS | 7 APPENDICES | JAVA 21 BASELINE | 8-WEEK STUDY PLAN", styles["cover_stat_line"]),
        Spacer(1, 24),
        Paragraph(EDITION, styles["cover_meta"]),
        Paragraph("Open educational printable interview study edition", styles["cover_meta"]),
        PageBreak(),
    ]


def copyright_story(styles: dict[str, ParagraphStyle]) -> list[Flowable]:
    paragraphs = [
        "Copyright and Disclaimer",
        f"Copyright 2026 {AUTHOR} and credited contributors. Open educational edition.",
        "Book prose, exercises, diagrams, and published editions are licensed under Creative Commons Attribution 4.0 International (CC BY 4.0). Build scripts and source code are licensed under MIT. Individual credit is recorded in AUTHORS.md and Git history.",
        "Repository: https://github.com/vinayreddykalluri/SDE2-Interview-Handbook. Attribution does not imply endorsement. Java and related marks are owned by their respective holders. Company names are used only to describe common interview markets.",
        "The Java platform, JVM implementations, tools, flags, support policies, and licensing terms evolve. Examples target Java 21 unless a section states otherwise. Verify release-specific and vendor-specific behavior against the primary sources in Appendix G before production use.",
        "The examples are provided without warranty. Review correctness, security, performance, operational, and legal requirements before adapting any example to a production system.",
        "Build provenance: generated from the ordered Markdown chapter sources in this project. The PDF and DOCX editions use the same content inventory.",
    ]
    title = Paragraph(paragraphs[0], styles["h1"])
    title._heading_level = 1
    title._bookmark_text = paragraphs[0]
    story: list[Flowable] = [Spacer(1, 80), title, HRFlowable(width="100%", thickness=1.2, color=GOLD, spaceAfter=18)]
    story.extend(Paragraph(html.escape(text), styles["copyright"]) for text in paragraphs[1:])
    story.append(PageBreak())
    return story


def toc_story(styles: dict[str, ParagraphStyle], fonts: dict[str, str]) -> list[Flowable]:
    toc_header = ParagraphStyle(
        "TOCRepeatHeader",
        fontName=fonts["bold"],
        fontSize=7.8,
        leading=9.5,
        textColor=colors.white,
        spaceAfter=0,
    )
    toc = GroupedTableOfContents(toc_header)
    toc.levelStyles = [
        ParagraphStyle("TOCPart", fontName=fonts["bold"], fontSize=9.8, leading=12.5, textColor=NAVY, leftIndent=0, firstLineIndent=0, spaceBefore=5),
        ParagraphStyle("TOCChapter", fontName=fonts["sans"], fontSize=8.5, leading=10.6, textColor=INK, leftIndent=18, firstLineIndent=0, rightIndent=4, spaceBefore=0.5),
    ]
    title = Paragraph("Contents", styles["toc_title"])
    title._heading_level = 1
    title._bookmark_text = "Contents"
    return [title, Paragraph("Parts, chapters, front matter, and appendices. PDF bookmarks also expose section-level navigation.", styles["small"]), Spacer(1, 8), toc]


def build_pdf(content_path: str) -> None:
    BUILD.mkdir(parents=True, exist_ok=True)
    pandoc = run(["pandoc", content_path, "--from=gfm+pipe_tables+task_lists", "--to=json"], capture=True)
    ast = json.loads(pandoc.stdout)
    (BUILD / "book-ast.json").write_text(json.dumps(ast), encoding="utf-8")
    fonts = register_fonts()
    styles = make_pdf_styles(fonts)
    renderer = PdfRenderer(styles, fonts)
    story: list[Flowable] = []
    story.extend(cover_story(styles, fonts))
    story.extend(copyright_story(styles))
    story.extend(toc_story(styles, fonts))
    story.extend(renderer.blocks(ast["blocks"]))
    doc = BookDocTemplate(str(FINAL_PDF), styles)
    doc.multiBuild(story)
    out = ROOT / "output" / "pdf"
    out.mkdir(parents=True, exist_ok=True)
    shutil.copy2(FINAL_PDF, out / FINAL_PDF.name)


def word_count() -> int:
    text = MASTER_MD.read_text(encoding="utf-8")
    text = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
    return len(re.findall(r"\b[\w'-]+\b", text))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--only", choices=("assemble", "docx", "pdf", "all"), default="all")
    parser.add_argument("--allow-incomplete", action="store_true")
    args = parser.parse_args()

    BUILD.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    DIST.mkdir(parents=True, exist_ok=True)
    FINAL_PDF.parent.mkdir(parents=True, exist_ok=True)
    content_path, docx_input = assemble_sources(args.allow_incomplete)
    if args.only in ("docx", "all"):
        build_docx(docx_input)
    if args.only in ("pdf", "all"):
        build_pdf(content_path)
    print(f"Markdown: {MASTER_MD}")
    print(f"Words: {word_count():,}")
    if FINAL_DOCX.exists():
        print(f"DOCX: {FINAL_DOCX}")
    if FINAL_PDF.exists():
        print(f"PDF: {FINAL_PDF}")


if __name__ == "__main__":
    main()
