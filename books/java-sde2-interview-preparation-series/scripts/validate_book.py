#!/usr/bin/env python3
"""Validate book sources, Java examples, and final publishing artifacts."""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
BOOK = ROOT / "content" / "master"
APPENDICES = BOOK / "appendices"

sys.path.insert(0, str(ROOT / "scripts"))
from build_book import APPENDIX_MAP, AUTHOR, CHAPTERS, FIGURES  # noqa: E402


REQUIRED_SECTIONS = [
    "Learning objectives",
    "Why this matters at SDE-2",
    "First-principles model",
    "Core terminology",
    "Detailed mechanics",
    "Worked Java example",
    "Execution or memory walkthrough",
    "Complexity and performance",
    "Edge cases and common mistakes",
    "Production engineering notes",
    "Interview questions and model answers",
    "Exercises",
    "Chapter summary",
    "Revision checklist",
]

# Chapters that teach through a narrative opening instead of the fixed template.
#
# REQUIRED_SECTIONS is a genuinely useful floor for reference chapters, but as a
# literal sequence it also mandates the shape "Learning objectives -> Why this
# matters -> Core terminology -> Detailed mechanics", which front-loads
# vocabulary before the reader has seen anything concrete. For the collections
# chapters that was the substance of a reader complaint: accurate material that
# read as definitions rather than explanation.
#
# So these chapters open with a specific situation and name the concept
# afterwards. The contract is not removed - it is restated as the part that
# carries the pedagogical guarantee: a worked example, a complexity treatment,
# and the full assessment tail. The narrative opening is free; everything a
# reader needs to practise and self-check is still mandatory.
#
# Membership is deliberately explicit so a chapter cannot drift out of the
# template by accident. Adding an id here should be an editorial decision.
NARRATIVE_CHAPTERS = frozenset({25, 26, 27, 28, 29, 30})

# The closing sequence every narrative chapter must still end on, in order.
NARRATIVE_TAIL = [
    "Edge cases and common mistakes",
    "Production engineering notes",
    "Interview questions and model answers",
    "Exercises",
    "Chapter summary",
    "Revision checklist",
]

# Prefixes that must each appear at least once before the tail.
NARRATIVE_REQUIRED_PREFIXES = ("Worked example", "Complexity")

MIN_NARRATIVE_TEACHING_SECTIONS = 4


def narrative_section_problems(h2: list[str]) -> list[str]:
    """Check a narrative chapter's structure, returning every problem found."""
    problems: list[str] = []
    if h2[-len(NARRATIVE_TAIL):] != NARRATIVE_TAIL:
        problems.append(
            "narrative chapter must end with, in order: "
            + " / ".join(NARRATIVE_TAIL)
        )
    body = h2[: -len(NARRATIVE_TAIL)] if len(h2) > len(NARRATIVE_TAIL) else []
    for prefix in NARRATIVE_REQUIRED_PREFIXES:
        if not any(section.startswith(prefix) for section in body):
            problems.append(f"narrative chapter has no {prefix!r} section")
    teaching = [
        section for section in body
        if not section.startswith(NARRATIVE_REQUIRED_PREFIXES)
    ]
    if len(teaching) < MIN_NARRATIVE_TEACHING_SECTIONS:
        problems.append(
            f"narrative chapter has {len(teaching)} teaching sections before the "
            f"worked example, expected at least {MIN_NARRATIVE_TEACHING_SECTIONS}"
        )
    return problems


FORBIDDEN = re.compile(
    r"\b(?:TODO|TBD|FIXME|placeholder)\b|This (?:chapter|appendix) is pending",
    re.IGNORECASE,
)
RUNNING_TEXT = {
    "java foundations to advanced engineering",
    "java 21 baseline",
    "appendices",
}


class Checks:
    def __init__(self) -> None:
        self.errors: list[str] = []
        self.warnings: list[str] = []

    def error(self, message: str) -> None:
        self.errors.append(message)

    def warn(self, message: str) -> None:
        self.warnings.append(message)


def substantive_characters(text: str) -> int:
    """Count page text after removing the generated running header/footer."""
    body: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        lowered = line.casefold()
        if not line or lowered in RUNNING_TEXT or line.isdigit():
            continue
        if re.fullmatch(r"Part [IVXLCDM]+ - .+", line):
            continue
        body.append(line)
    return len(re.sub(r"\s+", "", " ".join(body)))


def headings(text: str, level: int) -> list[str]:
    marker = "#" * level
    result: list[str] = []
    fenced = False
    for line in text.splitlines():
        if line.startswith("```"):
            fenced = not fenced
            continue
        if not fenced:
            match = re.match(rf"^{marker}\s+(.+)$", line)
            if match:
                result.append(match.group(1).strip())
    return result


def validate_chapters(checks: Checks) -> int:
    total_words = 0
    for number, title in CHAPTERS:
        paths = sorted(BOOK.glob(f"{number:02d}-*.md"))
        if len(paths) != 1:
            checks.error(f"Chapter {number}: expected one source, found {len(paths)}")
            continue
        path = paths[0]
        text = path.read_text(encoding="utf-8")
        total_words += len(re.findall(r"\b\w+[\w'-]*\b", text))

        h1 = headings(text, 1)
        source_title = ""
        if len(h1) == 1:
            source_title = re.sub(
                rf"^(?:Chapter\s+)?{number}\s*(?::|\.|-)\s*",
                "",
                h1[0],
                flags=re.IGNORECASE,
            ).strip()
        if source_title != title:
            checks.error(f"{path.name}: H1 title {source_title!r}, expected {title!r}")

        h2 = headings(text, 2)
        if number in NARRATIVE_CHAPTERS:
            for problem in narrative_section_problems(h2):
                checks.error(f"{path.name}: {problem}")
        elif h2 != REQUIRED_SECTIONS:
            checks.error(f"{path.name}: required H2 sequence differs ({len(h2)} headings)")

        if text.count("```") % 2:
            checks.error(f"{path.name}: unbalanced fenced code blocks")
        non_ascii = [(idx + 1, char) for idx, char in enumerate(text) if ord(char) > 127]
        if non_ascii:
            line = text[: non_ascii[0][0] - 1].count("\n") + 1
            checks.error(f"{path.name}:{line}: non-ASCII character {non_ascii[0][1]!r}")
        forbidden = FORBIDDEN.search(text)
        if forbidden:
            line = text[: forbidden.start()].count("\n") + 1
            checks.error(f"{path.name}:{line}: forbidden draft marker {forbidden.group(0)!r}")
        if "Specification boundary" not in text:
            checks.error(f"{path.name}: missing Specification boundary label")
        if number not in (50, 51, 52) and "HotSpot note" not in text:
            checks.error(f"{path.name}: missing HotSpot note label")
        word_count = len(re.findall(r"\b\w+[\w'-]*\b", text))
        if word_count < 1_700:
            checks.warn(f"{path.name}: short chapter ({word_count:,} words)")
    return total_words


def validate_front_and_appendices(checks: Checks) -> int:
    total_words = 0
    front_matter = (
        ("00-preface.md", "Preface"),
        ("00-about-the-author.md", "About the Author"),
        ("00-how-to-use-this-book.md", "How to Use This Book"),
        ("00-study-roadmap.md", "Study Roadmap"),
    )
    for filename, expected_h1 in front_matter:
        path = BOOK / filename
        if not path.exists():
            checks.error(f"Missing front matter: {filename}")
            continue
        text = path.read_text(encoding="utf-8")
        total_words += len(text.split())
        if headings(text, 1) != [expected_h1]:
            checks.error(f"{filename}: expected H1 {expected_h1!r}")
        if FORBIDDEN.search(text):
            checks.error(f"{filename}: contains draft marker")
        non_ascii = next((char for char in text if ord(char) > 127), None)
        if non_ascii is not None:
            checks.error(f"{filename}: contains non-ASCII character {non_ascii!r}")
        if filename == "00-about-the-author.md" and AUTHOR not in text:
            checks.error(f"{filename}: missing author name {AUTHOR!r}")

    for letter, title, filename in APPENDIX_MAP:
        path = APPENDICES / filename
        if not path.exists():
            checks.error(f"Missing Appendix {letter}: {filename}")
            continue
        text = path.read_text(encoding="utf-8")
        total_words += len(text.split())
        h1 = headings(text, 1)
        expected = f"Appendix {letter} - {title}"
        if h1 != [expected]:
            checks.error(f"{filename}: H1 {h1!r}, expected {expected!r}")
        if text.count("```") % 2:
            checks.error(f"{filename}: unbalanced fenced code blocks")
        if FORBIDDEN.search(text):
            checks.error(f"{filename}: contains draft marker")
        if len(text.split()) < 500:
            checks.warn(f"{filename}: short appendix ({len(text.split()):,} words)")
    return total_words


def validate_diagrams(checks: Checks) -> None:
    from PIL import Image

    for _, (filename, _) in FIGURES.items():
        path = ROOT / "assets" / "diagrams" / filename
        if not path.exists():
            checks.error(f"Missing diagram: {filename}")
            continue
        try:
            with Image.open(path) as image:
                width, height = image.size
                if width < 1_200 or height < 600:
                    checks.warn(f"{filename}: low dimensions {width}x{height}")
        except Exception as exc:  # pragma: no cover - defensive artifact check
            checks.error(f"{filename}: cannot open image: {exc}")


def validate_examples(checks: Checks) -> None:
    sources = sorted((ROOT / "examples" / "java" / "src" / "main" / "java").rglob("*.java"))
    if not sources:
        checks.error("No Java code examples found")
        return
    with tempfile.TemporaryDirectory(prefix="java-book-classes-") as directory:
        command = ["javac", "--release", "21", "-Xlint:all", "-d", directory, *map(str, sources)]
        try:
            result = subprocess.run(command, text=True, capture_output=True, timeout=120)
        except FileNotFoundError:
            checks.error("javac not found; JDK 21 is required")
            return
        if result.returncode:
            checks.error("Java examples failed to compile:\n" + result.stderr[-4_000:])
            return
        run = subprocess.run(
            ["java", "-cp", directory, "com.interviewbook.examples.AllExamplesSmokeTest"],
            text=True,
            capture_output=True,
            timeout=60,
        )
        if run.returncode or "passed" not in run.stdout:
            checks.error("Java example smoke test failed:\n" + (run.stdout + run.stderr)[-4_000:])


def validate_docx(checks: Checks) -> None:
    path = ROOT / "dist" / "java-sde2-interview-book.docx"
    if not path.exists():
        checks.error("Final DOCX is missing")
        return
    try:
        doc = Document(path)
    except Exception as exc:
        checks.error(f"DOCX cannot be opened: {exc}")
        return

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "About the Author" not in text or AUTHOR not in text:
        checks.error("DOCX is missing the About the Author front matter")
    if doc.core_properties.author != AUTHOR:
        checks.error(f"DOCX author metadata is {doc.core_properties.author!r}, expected {AUTHOR!r}")
    for number, title in CHAPTERS:
        if f"Chapter {number} - {title}" not in text:
            checks.error(f"DOCX missing Chapter {number} heading")
    for letter, title, _ in APPENDIX_MAP:
        if f"Appendix {letter} - {title}" not in text:
            checks.error(f"DOCX missing Appendix {letter} heading")
    if len(doc.inline_shapes) < len(FIGURES):
        checks.error(f"DOCX has {len(doc.inline_shapes)} inline figures; expected at least {len(FIGURES)}")
    if len(doc.tables) < 10:
        checks.warn(f"DOCX has only {len(doc.tables)} tables")
    for index, table in enumerate(doc.tables):
        cells = [cell.text.strip() for row in table.rows for cell in row.cells]
        if cells and sum(bool(value) for value in cells) / len(cells) < 0.5:
            checks.error(f"DOCX table {index + 1} is mostly empty")


def validate_pdf(checks: Checks) -> None:
    path = ROOT / "dist" / "java-sde2-interview-book.pdf"
    if not path.exists():
        checks.error("Final PDF is missing")
        return
    try:
        reader = PdfReader(path)
    except Exception as exc:
        checks.error(f"PDF cannot be opened: {exc}")
        return

    page_count = len(reader.pages)
    if not 450 <= page_count <= 700:
        checks.error(f"PDF page count {page_count} is outside requested 450-700 range")
    metadata = reader.metadata or {}
    if "Java Foundations" not in str(metadata.get("/Title", "")):
        checks.error("PDF title metadata is missing or incorrect")
    if str(metadata.get("/Author", "")) != AUTHOR:
        checks.error(f"PDF author metadata is {metadata.get('/Author')!r}, expected {AUTHOR!r}")

    page_texts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:
            checks.error(f"PDF page {index}: text extraction failed: {exc}")
            text = ""
        page_texts.append(text)
        if index > 1 and substantive_characters(text) < 8:
            checks.error(f"PDF page {index} appears blank")

    full_text = "\n".join(page_texts)
    cover_text = page_texts[0]
    if AUTHOR not in cover_text or "BY" not in cover_text:
        checks.error("PDF cover is missing the author name")
    for legacy_credit in ("FOUNDING AUTHOR", "EDITOR-IN-CHIEF", "CHIEF AUDITOR"):
        if legacy_credit in cover_text:
            checks.error(f"PDF cover still exposes the legacy credit: {legacy_credit}")
    if "About the Author" not in full_text or AUTHOR not in full_text:
        checks.error("PDF is missing the About the Author front matter")
    for number, title in CHAPTERS:
        if f"Chapter {number}" not in full_text or title not in full_text:
            checks.error(f"PDF missing Chapter {number} title")
    for letter, title, _ in APPENDIX_MAP:
        if f"Appendix {letter}" not in full_text or title not in full_text:
            checks.error(f"PDF missing Appendix {letter} title")


def validate_outputs(checks: Checks) -> None:
    for relative in (
        Path("dist/00-start-here/java-sde2-interview-book.md"),
        Path("dist/00-start-here/java-sde2-interview-book.pdf"),
        Path("dist/00-start-here/java-sde2-interview-book.docx"),
    ):
        path = ROOT / relative
        if not path.exists() or path.stat().st_size == 0:
            checks.error(f"Missing output: {relative}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-only", action="store_true", help="skip final PDF/DOCX checks")
    args = parser.parse_args()

    checks = Checks()
    chapter_words = validate_chapters(checks)
    other_words = validate_front_and_appendices(checks)
    validate_diagrams(checks)
    validate_examples(checks)
    if not args.source_only:
        validate_docx(checks)
        validate_pdf(checks)
        validate_outputs(checks)

    print(f"Validated {len(CHAPTERS)} chapters and {len(APPENDIX_MAP)} appendices")
    print(f"Approximate source words: {chapter_words + other_words:,}")
    for warning in checks.warnings:
        print(f"WARNING: {warning}")
    for error in checks.errors:
        print(f"ERROR: {error}")
    if checks.errors:
        print(f"FAILED with {len(checks.errors)} error(s)")
        return 1
    print("PASS")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
